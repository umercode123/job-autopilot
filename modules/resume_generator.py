# Job Autopilot - Resume Generator
# AI-powered resume tailoring with ATS optimization

import os
import json
from typing import Dict, List, Optional
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem, Image
from reportlab.lib.units import inch, mm
from reportlab.lib.colors import HexColor, white, black, lightgrey
from reportlab.graphics.shapes import Drawing, Line
from modules.ai_agent import ai_agent
from modules.logger_config import app_logger

load_dotenv()

class ResumeGenerator:
    """
    Generate ATS-optimized resumes tailored to job descriptions
    
    Features:
    - AI-powered keyword extraction from JD
    - Resume tailoring based on master resume + JD
    - .docx and PDF export
    - 1-page constraint
    - ATS-friendly formatting (no tables, simple layout)
    """
    
    def __init__(self):
        self.output_dir = "data/resumes"
        os.makedirs(self.output_dir, exist_ok=True)
        app_logger.info("Resume generator initialized")
    
    def load_master_resume(self, resume_path: str = None) -> Optional[Dict]:
        """
        Load master resume from markdown or JSON
        
        Args:
            resume_path: Path to master resume file
        
        Returns:
            dict: Parsed resume data
        """
        if not resume_path:
            # Try to find master resume in default locations
            possible_paths = [
                "Yuting Sun Master Resume.md",
                "data/master_resume.md",
                "master_resume.json"
            ]
            
            for path in possible_paths:
                if os.path.exists(path):
                    resume_path = path
                    break
        
        if not resume_path or not os.path.exists(resume_path):
            app_logger.error("Master resume not found!")
            return None
        
        try:
            with open(resume_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Parse markdown resume (simple parser)
            resume_data = self._parse_markdown_resume(content)
            app_logger.info(f"Loaded master resume from {resume_path}")
            return resume_data
        
        except Exception as e:
            app_logger.error(f"Failed to load master resume: {e}")
            return None
    
    def _parse_markdown_resume(self, content: str) -> Dict:
        """
        Parse markdown resume into structured data
        Supports various markdown heading levels (####, ##, #)
        """
        import re
        
        lines = content.split('\n')
        
        resume = {
            "name": "",
            "contact": {},
            "summary": "",
            "experience": [],
            "education": [],
            "skills": {},
            "certifications": [],
            "projects": []
        }
        
        current_section = None
        current_job = None
        summary_lines = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Skip empty lines
            if not line or line == '---' or line == '----':
                continue
            
            # Detect section headers (#### Level or ## Level)
            header_match = re.match(r'^#{1,6}\s+(.+)', line)
            if header_match:
                section_title = header_match.group(1).strip().lower()
                
                # Save previous job before switching sections
                if current_job and current_section == 'experience':
                    resume['experience'].append(current_job)
                    current_job = None
                
                if 'header' in section_title:
                    current_section = 'header'
                elif 'summary' in section_title or 'profile' in section_title:
                    current_section = 'summary'
                elif 'education' in section_title:
                    current_section = 'education'
                elif 'experience' in section_title:
                    current_section = 'experience'
                elif 'skill' in section_title:
                    current_section = 'skills'
                elif 'project' in section_title or 'saas' in section_title:
                    current_section = 'projects'
                continue
            
            # Parse Header section (name, contact)
            if current_section == 'header':
                # First line with | is contact info
                if '|' in line:
                    parts = [p.strip() for p in line.split('|')]
                    if len(parts) >= 4:
                        resume['name'] = parts[0]
                        resume['contact']['location'] = parts[1]
                        resume['contact']['phone'] = parts[2]
                        resume['contact']['email'] = parts[3]
                
                # Extract URLs
                if 'linkedin' in line.lower() or 'Linkedin' in line:
                    url = re.search(r'https?://[^\s]+', line)
                    if url:
                        resume['contact']['linkedin'] = url.group()
                
                if 'github' in line.lower():
                    url = re.search(r'https?://[^\s]+', line)
                    if url:
                        resume['contact']['github'] = url.group().rstrip(')')
                
                if 'portfolio' in line.lower() or 'Learning and Development' in line:
                    url = re.search(r'https?://[^\s]+', line)
                    if url:
                        resume['contact']['portfolio'] = url.group().rstrip(')')
            
            # Parse Summary
            elif current_section == 'summary':
                if line.startswith('*') and line.endswith('*'):
                    # Remove italic markers
                    summary_lines.append(line.strip('*').strip())
                elif line:
                    summary_lines.append(line)
            
            # Parse Education
            elif current_section == 'education':
                if line.startswith('-') or line.startswith('•'):
                    edu_line = line.lstrip('-•').strip()
                    
                    # Remove markdown bold markers
                    edu_line = edu_line.replace('**', '')
                    
                    # Split by em dash or double space
                    if '–' in edu_line:
                        parts = edu_line.split('–', 1)
                        resume['education'].append({
                            "title": parts[0].strip(),
                            "details": [parts[1].strip()] if len(parts) > 1 else []
                        })
                    else:
                        resume['education'].append({
                            "title": edu_line,
                            "details": []
                        })
            
            # Parse Experience
            elif current_section == 'experience':
                # Job title line (bold, starts with **)
                if line.startswith('**') and '—' in line:
                    # Save previous job
                    if current_job:
                        resume['experience'].append(current_job)
                    
                    # Parse: **Company — Job Title**
                    line_clean = line.strip('*').strip()
                    if '—' in line_clean:
                        parts = line_clean.split('—', 1)
                        company = parts[0].strip()
                        title = parts[1].strip()
                    else:
                        company = line_clean
                        title = ""
                    
                    current_job = {
                        "title": title,
                        "company": company,
                        "location": "",
                        "duration": "",
                        "details": []
                    }
                
                # Duration line (italic, starts with *) - Format: *Location | Duration*
                elif line.startswith('*') and ('|' in line or '–' in line):
                    duration_line = line.strip('*').strip()
                    if '|' in duration_line:
                        parts = duration_line.split('|')
                        if current_job:
                            # First part is location, last part is duration
                            current_job['location'] = parts[0].strip()
                            current_job['duration'] = parts[-1].strip()
                    else:
                        if current_job:
                            current_job['duration'] = duration_line
                
                # Bullet points
                elif line.startswith('-') or line.startswith('•'):
                    if current_job:
                        detail = line.lstrip('-•').strip()
                        current_job['details'].append(detail)
            
            # Parse Skills
            elif current_section == 'skills':
                if line.startswith('-') or line.startswith('•'):
                    skill_line = line.lstrip('-•').strip()
                    
                    # Check if it has a category (e.g., "EdTech Tools:")
                    if '**' in skill_line and ':' in skill_line:
                        # Extract category and skills
                        match = re.match(r'\*\*(.+?)\*\*:\s*(.+)', skill_line)
                        if match:
                            category = match.group(1).strip()
                            skills_content = match.group(2).strip()
                            resume['skills'][category] = skills_content
                    elif ':' in skill_line:
                        # Without bold markers
                        parts = skill_line.split(':', 1)
                        if len(parts) == 2:
                            category = parts[0].strip()
                            skills_content = parts[1].strip()
                            resume['skills'][category] = skills_content
                    else:
                        # Flat skill
                        if not isinstance(resume['skills'], dict):
                            resume['skills'] = []
                        if isinstance(resume['skills'], list):
                            resume['skills'].append(skill_line)
            
            # Parse Projects
            elif current_section == 'projects':
                if line.startswith('**') and ':' in line:
                    # Project title and description
                    project_line = line.strip('*').strip()
                    resume['projects'].append(project_line)
                elif line and not line.startswith('#'):
                    # Continue project description
                    if resume['projects']:
                        resume['projects'][-1] += ' ' + line
        
        # Add last job if exists
        if current_job and current_section == 'experience':
            resume['experience'].append(current_job)
        
        # Combine summary lines
        if summary_lines:
            resume['summary'] = ' '.join(summary_lines)
        
        app_logger.info(f"Parsed MD resume: {resume.get('name', 'Unknown')} - {len(resume.get('experience', []))} jobs")
        
        return resume
    
    def tailor_resume(self, master_resume: Dict, job_description: str, job_title: str, company: str) -> Dict:
        """
        Use AI to tailor resume for specific job
        
        Args:
            master_resume: Base resume data
            job_description: Target job JD
            job_title: Job title
            company: Company name
        
        Returns:
            dict: Tailored resume data
        """
        app_logger.info(f"Tailoring resume for {job_title} at {company}")
        
        # Build AI prompt
        prompt = f"""You are an expert resume writer specializing in ATS optimization.

**Task**: Tailor this resume for the following job, ensuring it passes ATS and appeals to HR.

**Job Details**:
- Title: {job_title}
- Company: {company}
- Job Description:
{job_description[:1500]}

**Master Resume**:
{json.dumps(master_resume, indent=2)}

**Requirements**:
1. Extract key skills/keywords from JD.
2. Reorder/emphasize relevant experience.
3. Add job-specific keywords naturally.
4. **IMPORTANT: If description is brief, EXPAND it using relevant industry keywords and standard responsibilities found in the JD.**
5. Keep 1-page length (approx 450-600 words).
6. Use action verbs and quantifiable achievements.
7. Return JSON format matching master resume structure.

**Output JSON** (experience items should include company, role, duration, bullets):
"""
        
        try:
            # Call GPT-4o-mini
            response = ai_agent.client.chat.completions.create(
                model=ai_agent.model,
                messages=[
                    {"role": "system", "content": "You are an expert ATS-optimized resume writer."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=2000
            )
            
            result = response.choices[0].message.content
            
            # Extract JSON from response
            if "```json" in result:
                result = result.split("```json")[1].split("```")[0].strip()
            elif "```" in result:
                result = result.split("```")[1].split("```")[0].strip()
            
            tailored_resume = json.loads(result)
            app_logger.info("Resume tailored successfully")
            
            return tailored_resume
        
        except Exception as e:
            app_logger.error(f"Resume tailoring failed: {e}")
            # Fallback: return master resume
            return master_resume
    
    def export_docx(self, resume_data: Dict, filename: str) -> str:
        """
        Export resume to .docx format (ATS-friendly)
        
        Args:
            resume_data: Resume data dict
            filename: Output filename
        
        Returns:
            str: Path to generated file
        """
        doc = Document()
        
        # Set narrow margins (ATS prefers simple layout)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Name (centered, bold, larger)
        name = doc.add_paragraph(resume_data.get('name', 'Your Name'))
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.runs[0].bold = True
        name.runs[0].font.size = Pt(16)
        
        # Contact info (centered, smaller)
        contact = resume_data.get('contact', {})
        contact_text = ' | '.join([
            contact.get('email', ''),
            contact.get('phone', ''),
            contact.get('location', ''),
            contact.get('linkedin', '')
        ])
        contact_para = doc.add_paragraph(contact_text)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()  # Spacer
        
        # Summary
        if resume_data.get('summary'):
            summary_heading = doc.add_paragraph('PROFESSIONAL SUMMARY')
            summary_heading.runs[0].bold = True
            summary_heading.runs[0].font.size = Pt(11)
            
            summary_text = doc.add_paragraph(resume_data['summary'])
            summary_text.runs[0].font.size = Pt(10)
            doc.add_paragraph()
        
        # Experience
        if resume_data.get('experience'):
            exp_heading = doc.add_paragraph('PROFESSIONAL EXPERIENCE')
            exp_heading.runs[0].bold = True
            exp_heading.runs[0].font.size = Pt(11)
            
            for exp in resume_data['experience']:
                # Job title + company
                job_line = doc.add_paragraph()
                job_line.add_run(exp.get('title', '')).bold = True
                
                if exp.get('company'):
                    job_line.add_run(f" | {exp['company']}")
                
                if exp.get('duration'):
                    job_line.add_run(f" | {exp['duration']}")
                
                job_line.runs[0].font.size = Pt(10)
                
                # Bullets
                for detail in exp.get('details', []):
                    bullet = doc.add_paragraph(detail, style='List Bullet')
                    bullet.runs[0].font.size = Pt(10)
                
                doc.add_paragraph()  # Spacer between jobs
        
        # Education
        if resume_data.get('education'):
            edu_heading = doc.add_paragraph('EDUCATION')
            edu_heading.runs[0].bold = True
            edu_heading.runs[0].font.size = Pt(11)
            
            for edu in resume_data['education']:
                edu_line = doc.add_paragraph(edu.get('title', ''))
                edu_line.runs[0].bold = True
                edu_line.runs[0].font.size = Pt(10)
                
                for detail in edu.get('details', []):
                    doc.add_paragraph(detail, style='List Bullet').runs[0].font.size = Pt(10)
        
        # Skills
        if resume_data.get('skills'):
            skills_heading = doc.add_paragraph('SKILLS')
            skills_heading.runs[0].bold = True
            skills_heading.runs[0].font.size = Pt(11)
            
            skills_text = ', '.join(resume_data['skills'])
            doc.add_paragraph(skills_text).runs[0].font.size = Pt(10)
        
        # Save
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        app_logger.info(f"Resume exported to {filepath}")
        
        return filepath
    
    
    def export_pdf(self, resume_data: Dict, filename: str) -> str:
        """
        Export resume to PDF with template-specific layout engine
        """
        filepath = os.path.join(self.output_dir, filename)
        
        # Determine layout engine
        template_name = resume_data.get('_meta', {}).get('template', 'classic_single_column')
        
        try:
            if "modern" in template_name:
                self._create_modern_layout(resume_data, filepath)
            else:
                self._create_classic_layout(resume_data, filepath)
                
            app_logger.info(f"PDF exported to {filepath} using {template_name} layout")
            return filepath
        except Exception as e:
            app_logger.error(f"PDF generation failed: {e}", exc_info=True)
            # Fallback to classic if modern fails
            try:
                self._create_classic_layout(resume_data, filepath)
                return filepath
            except:
                return ""

    def _create_classic_layout(self, resume_data: Dict, filepath: str):
        """
        Classic Serif Layout (Resume Matcher Style)
        - Times New Roman
        - Horizontal lines under headers
        - Right-aligned dates
        """
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            topMargin=0.5*inch, bottomMargin=0.5*inch,
            leftMargin=0.5*inch, rightMargin=0.5*inch
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Fonts
        title_font = 'Times-Bold'
        body_font = 'Times-Roman'
        
        # Styles
        style_name = ParagraphStyle(
            'Name', 
            parent=styles['Normal'], 
            fontName=title_font, 
            fontSize=24, 
            leading=28,  # Critical fix for overlap
            alignment=1, 
            spaceAfter=12 # Critical fix for overlap
        )
        style_role = ParagraphStyle('Role', parent=styles['Normal'], fontName=body_font, fontSize=11, alignment=1, textColor=HexColor('#555555'), spaceAfter=6)
        style_contact = ParagraphStyle(
            'Contact', 
            parent=styles['Normal'], 
            fontName=body_font, 
            fontSize=10, 
            alignment=1, 
            spaceAfter=12
        )
        
        style_section = ParagraphStyle('Section', parent=styles['Normal'], fontName=title_font, fontSize=12, spaceBefore=12, spaceAfter=4, textTransform='uppercase')
        style_job_title = ParagraphStyle('JobTitle', parent=styles['Normal'], fontName='Times-Bold', fontSize=10.5)
        style_job_meta = ParagraphStyle('JobMeta', parent=styles['Normal'], fontName=body_font, fontSize=10.5, alignment=2) # Right align
        style_bullet = ParagraphStyle('Bullet', parent=styles['Normal'], fontName=body_font, fontSize=10, leading=13, leftIndent=12, bulletIndent=0)
        
        def draw_line(width):
            d = Drawing(width, 1)
            d.add(Line(0, 0, width, 0))
            return d

        # --- HEADER ---
        story.append(Paragraph(resume_data.get('name', 'YOUR NAME').upper(), style_name))
        
        # Contact Line
        contact = resume_data.get('contact', {})
        parts = []
        if contact.get('email'): parts.append(contact['email'])
        if contact.get('phone'): parts.append(contact['phone'])
        if contact.get('location'): parts.append(contact['location'])
        if contact.get('linkedin'): parts.append(f'<a href="{contact["linkedin"]}">LinkedIn</a>')
        if contact.get('portfolio'): parts.append(f'<a href="{contact["portfolio"]}">Portfolio</a>')
        if contact.get('github'): parts.append(f'<a href="{contact["github"]}">Other</a>')
        
        story.append(Paragraph(" • ".join(parts), style_contact))
        
        # Line Divider
        story.append(draw_line(500))
        story.append(Spacer(1, 10))

        # --- SECTIONS ---
        order = resume_data.get('_meta', {}).get('section_order', ['summary', 'experience', 'projects', 'skills', 'education'])
        
        for section in order:
            if section == 'summary' and resume_data.get('summary'):
                story.append(Paragraph("SUMMARY", style_section))
                # Removed line under Summary as requested
                story.append(Spacer(1, 4))
                story.append(Paragraph(resume_data['summary'], ParagraphStyle('Sum', parent=style_bullet, leftIndent=0)))
            
            elif section == 'experience' and resume_data.get('experience'):
                story.append(Paragraph("EXPERIENCE", style_section))
                story.append(draw_line(500))
                story.append(Spacer(1, 6))
                
                for exp in resume_data['experience']:
                    # Job Header Table (Left: Title/Company, Right: Date/Location)
                    # Row 1: Company (Left) | Date (Right)
                    # Row 2: Title (Left)   | Location (Right)
                    
                    c1 = [Paragraph(f"<b>{exp.get('company','')}</b>", style_job_title)]
                    c2 = [Paragraph(exp.get('duration',''), style_job_meta)]
                    
                    t1 = Table([[c1, c2]], colWidths=[4.5*inch, 2.5*inch])
                    t1.setStyle(TableStyle([('ALIGN', (1,0), (1,0), 'RIGHT'), ('VALIGN', (0,0), (-1,-1), 'TOP')]))
                    story.append(t1)
                    
                    c3 = [Paragraph(f"<i>{exp.get('title','')}</i>", ParagraphStyle('Italic', parent=style_job_title, fontName='Times-Italic'))]
                    c4 = [Paragraph(exp.get('location', ''), style_job_meta)]  # FIXED: Use job-specific location, not contact
                    t2 = Table([[c3, c4]], colWidths=[5.5*inch, 1.5*inch]) 
                    # FIX: Append t2 (Job Title row) which was missing
                    story.append(t2) 
                    story.append(Spacer(1, 1))

                    for detail in exp.get('details', []):
                        story.append(Paragraph(f"• {detail}", style_bullet))
                    story.append(Spacer(1, 8))

            elif section == 'skills' and resume_data.get('skills'):
                story.append(Paragraph("SKILLS", style_section))
                story.append(draw_line(500))
                story.append(Spacer(1, 4))
                
                # Render skills line by line
                skills_data = resume_data['skills']
                
                # Ensure it is a list
                if isinstance(skills_data, str):
                     skills_data = [skills_data]
                elif isinstance(skills_data, dict):
                     # Fallback if somehow dict passed (should be converted by UI now)
                     skills_data = [f"{k}: {v}" for k,v in skills_data.items()]
                     
                for skill_line in skills_data:
                    # Format: "Category: Skill, Skill" -> "<b>Category</b>: Skill, Skill"
                    text = str(skill_line)
                    if ":" in text:
                        parts = text.split(":", 1)
                        formatted_text = f"<b>{parts[0]}</b>:{parts[1]}"
                    else:
                        formatted_text = text
                        
                    story.append(Paragraph(formatted_text, style_bullet))

                        
            elif section == 'projects' and resume_data.get('projects'):
                story.append(Paragraph("PROJECTS", style_section))
                story.append(draw_line(500))
                story.append(Spacer(1, 4))
                
                # Helper function to clean markdown syntax
                def clean_markdown(text):
                    """Remove markdown syntax like **bold** and keep plain text"""
                    import re
                    # Remove **bold** -> bold
                    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                    text = re.sub(r'\*\*(.+)\*\*', r'\1', text)
                    # Remove leftover ** at end of strings
                    text = re.sub(r'\*\*$', '', text)
                    text = re.sub(r'^\*\*', '', text)
                    # Remove *italic* -> italic
                    text = re.sub(r'\*(.+?)\*', r'\1', text)
                    # Remove __bold__ -> bold
                    text = re.sub(r'__(.+?)__', r'\1', text)
                    # Remove any remaining standalone * or **
                    text = text.replace('**', '').replace('*', '')
                    return text.strip()
                
                # Projects Logic (Handle both List[str] and List[Dict])
                projects = resume_data['projects']
                for proj in projects:
                    if isinstance(proj, dict):
                        # Structured Project
                        p_title = f"<b>{proj.get('title', 'Project')}</b>"
                        if proj.get('link'):
                            p_title += f" | <a href='{proj['link']}'>Link</a>"
                        
                        # Row: Title (Left) | Tech Stack or Date (Right)
                        c1 = [Paragraph(p_title, style_bullet)]
                        c2 = [Paragraph(proj.get('tech_stack', ''), style_job_meta)]
                        t = Table([[c1, c2]], colWidths=[5*inch, 2*inch])
                        story.append(t)
                        
                        if proj.get('details'):
                            for d in proj['details']:
                                clean_d = clean_markdown(str(d))
                                story.append(Paragraph(f"• {clean_d}", style_bullet))
                    else:
                        # Simple String Project - Clean markdown
                        clean_proj = clean_markdown(str(proj))
                        story.append(Paragraph(f"• {clean_proj}", style_bullet))
                    
                    story.append(Spacer(1, 4))

            elif section == 'education' and resume_data.get('education'):
                story.append(Paragraph("EDUCATION", style_section))
                story.append(draw_line(500))
                story.append(Spacer(1, 6))
                
                for edu in resume_data['education']:
                    title_line = f"<b>{edu.get('title','')}</b>"
                    # Append Institution and Date if available, e.g. "Master of X | University, 2024"
                    # But often title has degree, details has uni. Let's keep it simple or strictly follow JSON.
                    # Current JSON structure usually puts Uni in 'details'.
                    
                    c1 = Paragraph(title_line, style_job_title)
                    story.append(c1)
                    if edu.get('details'):
                        for d in edu['details']:
                             story.append(Paragraph(d, style_bullet))
                    story.append(Spacer(1, 4))
        
        doc.build(story)

    def _create_modern_layout(self, resume_data: Dict, filepath: str):
        """
        Modern Sidebar Layout
        - Helvetica/Sans-Serif
        - Dark Sidebar (Left) with Contact & Skills
        - Light Main Content (Right) with Summary & Experience
        """
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            topMargin=0, bottomMargin=0, leftMargin=0, rightMargin=0 # Handled by Table
        )
        
        styles = getSampleStyleSheet()
        
        # Colors
        bg_color = HexColor('#2C3E50') # Midnight Blue
        text_white = HexColor('#FFFFFF')
        text_dark = HexColor('#2c3e50')
        
        # Styles
        style_sidebar_text = ParagraphStyle('Side', fontName='Helvetica', fontSize=9, textColor=text_white, leading=14)
        style_sidebar_header = ParagraphStyle('SideHead', fontName='Helvetica-Bold', fontSize=11, textColor=text_white, spaceBefore=10, spaceAfter=4)
        style_main_header = ParagraphStyle('MainHead', fontName='Helvetica-Bold', fontSize=22, textColor=text_dark, spaceAfter=8)
        style_role = ParagraphStyle('Role', fontName='Helvetica', fontSize=12, textColor=HexColor('#7f8c8d'), spaceAfter=18)
        style_section = ParagraphStyle('Sec', fontName='Helvetica-Bold', fontSize=12, textColor=text_dark, spaceBefore=12, spaceAfter=6, textTransform='uppercase')
        style_job = ParagraphStyle('Job', fontName='Helvetica-Bold', fontSize=11, textColor=black)
        style_bullet = ParagraphStyle('Bull', fontName='Helvetica', fontSize=9.5, leading=13, leftIndent=10, textColor=black)
        
        # --- LEFT COLUMN (SIDEBAR) ---
        sidebar_content = []
        sidebar_content.append(Spacer(1, 0.5*inch))
        
        # Contact
        sidebar_content.append(Paragraph("CONTACT", style_sidebar_header))
        contact = resume_data.get('contact', {})
        if contact.get('email'): sidebar_content.append(Paragraph(f"📧 {contact['email']}", style_sidebar_text))
        if contact.get('phone'): sidebar_content.append(Paragraph(f"📱 {contact['phone']}", style_sidebar_text))
        if contact.get('location'): sidebar_content.append(Paragraph(f"📍 {contact['location']}", style_sidebar_text))
        sidebar_content.append(Spacer(1, 0.2*inch))
        
        # Skills (Move to sidebar in Modern)
        if resume_data.get('skills'):
            sidebar_content.append(Paragraph("SKILLS", style_sidebar_header))
            skills = resume_data['skills']
            if isinstance(skills, dict):
                for cat, val in skills.items():
                    sidebar_content.append(Paragraph(f"<b>{cat}</b>", style_sidebar_text))
                    sidebar_content.append(Paragraph(val, style_sidebar_text))
                    sidebar_content.append(Spacer(1, 0.1*inch))
            else:
                for s in skills[:15]: # Limit
                    sidebar_content.append(Paragraph(f"• {s}", style_sidebar_text))
        
        def draw_line(width, color=lightgrey):
            d = Drawing(width, 1)
            d.add(Line(0, 0, width, 0, strokeColor=color))
            return d

        # --- RIGHT COLUMN (MAIN) ---
        main_content = []
        main_content.append(Spacer(1, 0.5*inch))
        
        # Name
        main_content.append(Paragraph(resume_data.get('name', 'Name'), style_main_header))
        # Summary
        if resume_data.get('summary'):
            main_content.append(Paragraph("PROFILE", style_section))
            main_content.append(Paragraph(resume_data['summary'], style_bullet))
        
        # Experience
        if resume_data.get('experience'):
            main_content.append(Paragraph("EXPERIENCE", style_section))
            main_content.append(draw_line(400))
            for exp in resume_data['experience']:
                main_content.append(Spacer(1, 8))
                main_content.append(Paragraph(f"{exp.get('title')} at {exp.get('company')}", style_job))
                main_content.append(Paragraph(f"<font color='#7f8c8d'>{exp.get('duration')}</font>", style_bullet))
                for d in exp.get('details', []):
                    main_content.append(Paragraph(f"• {d}", style_bullet))
        
        # Education
        if resume_data.get('education'):
            main_content.append(Paragraph("EDUCATION", style_section))
            for edu in resume_data['education']:
                main_content.append(Paragraph(f"<b>{edu.get('title')}</b>", style_bullet))

        # Main Layout Table
        # Using a fixed high height to simulate sidebar filling page (hacky but works for 1 page)
        # Better: calculate height or use Frames. For now, assuming ~10 inch content.
        
        tbl = Table([[sidebar_content, main_content]], colWidths=[2.3*inch, 5.7*inch])
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (0,0), bg_color),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('LEFTPADDING', (0,0), (0,0), 15),
            ('RIGHTPADDING', (0,0), (0,0), 10),
            ('LEFTPADDING', (1,0), (1,0), 20),
            ('RIGHTPADDING', (1,0), (1,0), 20),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
        ]))
        
        doc.build([tbl])
    
    # ============================================================
    # Template System (Phase 2 - Resume Export Feature)
    # ============================================================
    
    def load_template(self, template_name: str) -> Optional[Dict]:
        """
        Load template configuration from JSON
        
        Args:
            template_name: Template name (e.g., "classic_single_column")
        
        Returns:
            dict: Template configuration
        """
        template_path = f"data/templates/{template_name}.json"
        
        if not os.path.exists(template_path):
            app_logger.error(f"Template not found: {template_name}")
            return None
        
        try:
            with open(template_path, 'r', encoding='utf-8') as f:
                template = json.load(f)
            app_logger.info(f"Loaded template: {template_name}")
            return template
        except Exception as e:
            app_logger.error(f"Failed to load template: {e}")
            return None
    
    def apply_template(self, resume_data: Dict, template: Dict) -> Dict:
        """
        Apply template settings to resume data
        
        Args:
            resume_data: Resume content
            template: Template configuration
        
        Returns:
            dict: Resume with template metadata
        """
        if not template:
            app_logger.warning("apply_template received None template, returning original data")
            return resume_data

        # Reorder sections based on template
        ordered_resume = {}
        for section in template['section_order']:
            if section in resume_data:
                ordered_resume[section] = resume_data[section]
        
        # Keep other sections that aren't in template order
        for key, value in resume_data.items():
            if key not in ordered_resume:
                ordered_resume[key] = value
        
        # Add template metadata
        ordered_resume['_meta'] = {
            'template': template['name'],
            'section_order': template['section_order'], # IMPORTANT: Persist section order for PDF generation
            'line_spacing': template.get('line_spacing', 1.0),
            'fonts': template.get('fonts', {}),
            'margins': template.get('margins', {}),
            'layout': template.get('layout', 'single_column')
        }
        
        app_logger.info(f"Applied template: {template['name']}")
        return ordered_resume
    
    # ============================================================
    # Multi-Format Resume Parsing
    # ============================================================
    
    def parse_pdf_resume(self, pdf_path: str) -> Optional[Dict]:
        """
        Parse PDF resume using pdfminer.six (better than PyPDF2) + AI structuring
        
        Args:
            pdf_path: Path to PDF file
        
        Returns:
            dict: Structured resume data
        """
        try:
            from pdfminer.high_level import extract_text
            
            app_logger.info(f"Parsing PDF resume: {pdf_path}")
            
            # Extract text from PDF (pdfminer.six preserves more structure)
            text = extract_text(pdf_path)
            
            app_logger.info(f"Extracted {len(text)} characters from PDF")
            
            # If text is too short, might be a scanned PDF
            if len(text) < 100:
                app_logger.warning("PDF text too short, might be scanned. Consider OCR.")
                return None
            
            # Use AI to structure the text
            return self._parse_with_ai(text)
        
        except Exception as e:
            app_logger.error(f"Failed to parse PDF: {e}")
            return None
    
    def parse_docx_resume(self, docx_path: str) -> Optional[Dict]:
        """
        Parse DOCX resume using docx2txt (better text extraction than python-docx)
        
        Args:
            docx_path: Path to DOCX file
        
        Returns:
            dict: Structured resume data
        """
        try:
            import docx2txt
            
            app_logger.info(f"Parsing DOCX resume: {docx_path}")
            
            # Extract all text including text from tables
            text = docx2txt.process(docx_path)
            
            app_logger.info(f"Extracted {len(text)} characters from DOCX")
            
            if len(text) < 100:
                app_logger.warning("DOCX text too short")
                return None
            
            # Use AI to structure the text
            return self._parse_with_ai(text)
        
        except Exception as e:
            app_logger.error(f"Failed to parse DOCX: {e}")
            return None
    
    def _parse_with_ai(self, text: str) -> Optional[Dict]:
        """
        Use AI to parse unstructured resume text into JSON
        CRITICAL: Preserve ALL details, especially:
        - Each job's individual bullet points
        - Skills categorization
        - All links and contact info
        
        Args:
            text: Raw resume text
        
        Returns:
            dict: Structured resume data
        """
        try:
            # Use full text, don't truncate
            text_to_parse = text[:5000] if len(text) > 5000 else text
            
            prompt = f"""Parse this resume into JSON. EXTRACT ALL INFORMATION WITHOUT LOSING ANY DETAILS.

Resume Text:
{text_to_parse}

CRITICAL RULES:
1. **PRESERVE EVERY BULLET POINT** under each job separately
2. **KEEP SKILLS CATEGORIZED** if they have categories (e.g., "EdTech Tools:", "AI & Development:")
3. **EXTRACT ALL URLs** (LinkedIn, GitHub, Portfolio)
4. **KEEP ALL JOB DETAILS** (each job must have its own details array)

Return EXACTLY this JSON structure:
{{
  "name": "Full Name",
  "contact": {{
    "email": "email@example.com",
    "phone": "phone number",
    "location": "city, province",
    "linkedin": "COMPLETE LinkedIn URL",
    "github": "COMPLETE GitHub URL", 
    "portfolio": "COMPLETE Portfolio URL"
  }},
  "summary": "complete professional summary",
  "experience": [
    {{
      "title": "Job Title 1",
      "company": "Company Name",
      "duration": "Start Date - End Date",
      "details": [
        "EVERY bullet point for THIS job",
        "Do NOT merge with other jobs",
        "Keep ALL achievements"
      ]
    }},
    {{
      "title": "Job Title 2",
      "company": "Company 2",
      "duration": "Start - End",
      "details": [
        "All bullet points for THIS job only"
      ]
    }}
  ],
  "education": [
    {{
      "title": "Degree, Major",
      "details": ["University, Dates"]
    }}
  ],
  "skills": {{
    "EdTech Tools and Instructional Design": "H5P, Articulate Suite, Canva, etc",
    "AI & Development": "RAG, Python, JavaScript, etc",
    "HR Practices": "Recruitment, etc"
  }},
  "projects": ["project if present"],
  "certifications": ["cert if present"]
}}

IMPORTANT:
- If skills have categories (e.g., "EdTech Tools:", "AI & Development:"), use dict format
- If skills are flat list, return array
- DO NOT merge bullet points from different jobs
- KEEP EVERY SINGLE DETAIL

Return ONLY valid JSON, no markdown, no extra text."""

            response = ai_agent.client.chat.completions.create(
                model=ai_agent.model,
                messages=[
                    {"role": "system", "content": "You are a resume parser. Extract ALL information exactly as written. Do NOT summarize or merge content. PRESERVE structure. Return ONLY valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.05,  # Extremely low for accurate extraction
                max_tokens=3000
            )
            
            result = response.choices[0].message.content
            
            # Extract JSON from response
            if "```json" in result:
                result = result.split("```json")[1].split("```")[0].strip()
            elif "```" in result:
                result = result.split("```")[1].split("```")[0].strip()
            
            resume_data = json.loads(result)
            
            # Log what was extracted
            job_count = len(resume_data.get('experience', []))
            app_logger.info(f"Successfully parsed resume: {resume_data.get('name', 'Unknown')} - {job_count} jobs")
            
            return resume_data
        
        except Exception as e:
            app_logger.error(f"AI parsing failed: {e}")
            return None
    
    # ============================================================
    # AI Compression (1-page constraint)
    # ============================================================
    
    def tailor_and_compress(self, resume_data: Dict, job_description: str, 
                            template: Optional[Dict] = None, 
                            mode: str = "balanced") -> Optional[Dict]:
        """
        Tailor resume for specific job AND compress to fit 1 page using AI
        
        Args:
            resume_data: Full resume
            job_description: Target job description
            template: Template config (for word limit)
            mode: 'aggressive', 'balanced', or 'conservative'
        
        Returns:
            dict: Tailored and compressed resume
        """
        try:
            # Compression/Optimization settings
            # We want a STRICT ONE-PAGE resume.
            # 420 words is safer for 1 page with projects section
            
            target_words = 350  # Very aggressive for strict 1-page with projects
            if template and template.get('layout') == 'two_column':
                target_words = 420
            
            app_logger.info(f"Tailoring resume. Target: ~{target_words} words ({mode} mode)")
            
            # AI Prompt - STRICT 1-PAGE & KEYWORDS + ANTI-HALLUCINATION
            prompt = f"""You are an expert resume writer creating a STRICT 1-PAGE resume.

TARGET JOB DESCRIPTION:
{job_description[:2500]}

CURRENT RESUME (JSON):
{json.dumps(resume_data, indent=2)}

TASK: Create a 1-PAGE optimized resume using JD keywords. You MUST include ALL sections.

CRITICAL 1-PAGE CONSTRAINTS:
- Total resume: ~{target_words} words maximum (STRICT LIMIT)
- Experience: MAX 2-3 bullets per job (ONLY the most relevant to JD)
- Each bullet: MAX 1 line only
- Projects: MAX 1 short sentence per project
- Skills: ONLY keep skills mentioned in JD, remove irrelevant skills
- Education: One line per degree only

MANDATORY SECTIONS (DO NOT OMIT ANY):
1. Summary (2 sentences max)
2. Experience (all jobs with compressed bullets, keep location field!)
3. Education
4. Skills (filtered to match JD keywords)
5. **PROJECTS** - Include but compress heavily

RULES:
1. **KEYWORD INTEGRATION**: Rewrite bullets using JD keywords.
2. **STRUCTURE**: Keep exact JSON keys including 'projects' and 'location'.
3. **NO FILLER**: Remove wordy phrases. Start with Action Verbs.
4. **SKILLS FILTER**: Only keep skills that appear in the JD or are closely related.

CRITICAL DATA PRESERVATION (DO NOT CHANGE THESE):
- Copy 'location' field EXACTLY from each job in source JSON
- Copy 'duration' field EXACTLY from each job in source JSON  
- Copy 'company' field EXACTLY from each job in source JSON
- Copy 'title' field EXACTLY from each job in source JSON

PROJECTS HANDLING:
- Keep ALL 3 project titles
- Each project description: 1 short sentence only (no markdown syntax)
- Remove ** and * markdown characters

OUTPUT JSON MUST INCLUDE:
- "name", "contact", "summary", "experience" (with location!), "education", "skills", "projects"

Return JSON:"""

            response = ai_agent.client.chat.completions.create(
                model=ai_agent.model,
                messages=[
                    {"role": "system", "content": "You are a specialized resume writer. Your goal is to create a DENSE, KEYWORD-RICH resume that passes ATS and fills the page. NEVER change locations, dates, or company names. Return ONLY valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,  # Reduced from 0.7 to minimize hallucination
                max_tokens=3500
            )
            
            result = response.choices[0].message.content
            
            # Extract JSON
            if "```json" in result:
                result = result.split("```json")[1].split("```")[0].strip()
            elif "```" in result:
                result = result.split("```")[1].split("```")[0].strip()
            
            compressed = json.loads(result)
            
            # ============================================================
            # POST-PROCESSING: CRITICAL - Restore fields AI might change
            # Use COMPANY NAME MATCHING instead of index (AI may reorder)
            # ============================================================
            
            original_exp = resume_data.get('experience', [])
            compressed_exp = compressed.get('experience', [])
            
            # Create lookup by company name (case-insensitive)
            original_lookup = {}
            for exp in original_exp:
                company = exp.get('company', '').strip().lower()
                if company:
                    original_lookup[company] = exp
            
            # Restore critical fields for each job using company name matching
            restored_count = 0
            for comp_exp in compressed_exp:
                comp_company = comp_exp.get('company', '').strip().lower()
                
                # Find matching original job
                orig = original_lookup.get(comp_company)
                if not orig:
                    # Try partial match (e.g., "Parking Control" vs "Parking Control Services Group")
                    for key, val in original_lookup.items():
                        if key in comp_company or comp_company in key:
                            orig = val
                            break
                
                if orig:
                    # Force restore location - THIS IS CRITICAL
                    if orig.get('location'):
                        comp_exp['location'] = orig['location']
                    # Force restore company name (exact spelling)
                    if orig.get('company'):
                        comp_exp['company'] = orig['company']
                    # Force restore duration/dates
                    if orig.get('duration'):
                        comp_exp['duration'] = orig['duration']
                    restored_count += 1
                    app_logger.debug(f"Restored fields for: {comp_exp.get('company')}")
                else:
                    app_logger.warning(f"Could not match job: {comp_exp.get('company')} - fields may be incorrect")
            
            app_logger.info(f"Post-processing: Restored {restored_count}/{len(compressed_exp)} jobs")
            
            # 2. ALWAYS restore projects section from original
            if resume_data.get('projects'):
                compressed['projects'] = resume_data['projects']
                app_logger.info("Projects section restored from original")
            
            # 3. Restore contact info (should never change)
            if resume_data.get('contact'):
                compressed['contact'] = resume_data['contact']
            
            # 4. Restore education (should rarely change)
            if resume_data.get('education'):
                compressed['education'] = resume_data['education']
            
            # 5. Preserve _meta from original (section order, template, etc.)
            if resume_data.get('_meta'):
                compressed['_meta'] = resume_data['_meta']
            
            # Add word count metadata
            word_count = self._count_words(compressed)
            compressed['_word_count'] = word_count
            
            app_logger.info(f"Tailored resume generated: {word_count} words (with strict post-processing)")
            
            return compressed
        
        except Exception as e:
            app_logger.error(f"Tailoring and compression failed: {e}")
            return None
    
    def _count_words(self, resume_data: Dict) -> int:
        """Count approximate words in resume"""
        text = json.dumps(resume_data)
        # Remove JSON syntax
        text = text.replace('{', '').replace('}', '').replace('[', '').replace(']', '')
        text = text.replace('"', '').replace(',', ' ')
        words = text.split()
        return len(words)

# Global instance
resume_generator = ResumeGenerator()

if __name__ == "__main__":
    # Test
    gen = ResumeGenerator()
    master = gen.load_master_resume()
    
    if master:
        print("Master resume loaded:")
        print(f"Name: {master.get('name')}")
        print(f"Experience items: {len(master.get('experience', []))}")
